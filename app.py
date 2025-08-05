import pandas as pd
import json
import os
from datetime import datetime
import PyPDF2
import docx
from io import BytesIO
import glob
from pathlib import Path
import re
from email_validator import validate_email, EmailNotValidError
import streamlit as st



# Enhanced rule-based resume parser with improved accuracy
class ResumeParser:
    def __init__(self):
        # Enhanced email regex pattern
        self.email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b'
        
        # Enhanced phone regex patterns
        self.phone_patterns = [
            r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US format
            r'(\+\d{1,3}[-.\s]?)?\d{10,15}',  # International
            r'(\+\d{1,3}[-.\s]?)?\d{4}[-.\s]?\d{3}[-.\s]?\d{3}',  # Different format
            r'(\+\d{1,3}[-.\s]?)?\d{3}[-.\s]?\d{4}[-.\s]?\d{4}',  # Another format
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # Simple 10-digit
        ]
        
        # Expanded and categorized skills
        self.technical_skills = [
            # Programming Languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust',
            'swift', 'kotlin', 'scala', 'r', 'matlab', 'perl', 'shell scripting', 'bash',
            
            # Web Technologies
            'html', 'css', 'react', 'angular', 'vue.js', 'node.js', 'express.js', 'django', 'flask',
            'spring boot', 'laravel', 'asp.net', 'jquery', 'bootstrap', 'sass', 'less',
            
            # Databases
            'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'sqlite', 'redis', 'cassandra',
            'dynamodb', 'elasticsearch', 'neo4j',
            
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'gitlab', 'github',
            'terraform', 'ansible', 'chef', 'puppet', 'ci/cd', 'microservices',
            
            # Data Science & AI
            'machine learning', 'deep learning', 'data science', 'artificial intelligence',
            'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy', 'matplotlib', 'seaborn',
            'tableau', 'power bi', 'excel', 'spss', 'sas', 'hadoop', 'spark',
            
            # Design & Creative
            'photoshop', 'illustrator', 'figma', 'sketch', 'adobe creative suite', 'canva',
            'ui/ux design', 'graphic design', 'web design',
            
            # Project Management
            'agile', 'scrum', 'kanban', 'jira', 'trello', 'asana', 'project management'
        ]
        
        # Enhanced education patterns
        self.education_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'degree', 'university', 'college',
            'institute', 'school', 'graduation', 'graduate', 'undergraduate', 'mba', 'ms', 'ma',
            'btech', 'be', 'bsc', 'msc', 'ba', 'bs', 'diploma', 'certification', 'associate',
            'bca', 'mca', 'bba', 'engineering', 'computer science', 'information technology'
        ]
        
        # Enhanced experience keywords
        self.experience_keywords = [
            'experience', 'worked', 'employed', 'position', 'role', 'job', 'career',
            'professional', 'intern', 'internship', 'volunteer', 'freelance', 'consultant',
            'manager', 'developer', 'engineer', 'analyst', 'coordinator', 'specialist',
            'associate', 'senior', 'junior', 'lead', 'director', 'executive'
        ]
        
        # Language keywords
        self.language_keywords = [
            'english', 'spanish', 'french', 'german', 'chinese', 'japanese', 'korean',
            'hindi', 'arabic', 'portuguese', 'russian', 'italian', 'dutch', 'swedish',
            'fluent', 'native', 'conversational', 'basic', 'intermediate', 'advanced'
        ]
        
        # Certification keywords
        self.certification_keywords = [
            'certified', 'certification', 'certificate', 'license', 'credential',
            'aws certified', 'microsoft certified', 'google certified', 'cisco certified',
            'pmp', 'scrum master', 'agile', 'itil', 'comptia', 'cissp'
        ]

    def extract_email(self, text):
        emails = re.findall(self.email_pattern, text, re.IGNORECASE)
        # Remove duplicates and filter out common false positives
        valid_emails = []
        for email in emails:
            if '@' in email and '.' in email.split('@')[1]:
                valid_emails.append(email.lower())
        
        return valid_emails[0] if valid_emails else "Not specified"

    def extract_phone(self, text):
        phones = []
        for pattern in self.phone_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    phone = ''.join(match)
                else:
                    phone = match
                
                # Clean and validate phone number
                phone = re.sub(r'[^\d+]', '', phone)
                if len(phone.replace('+', '')) >= 10:
                    phones.append(phone)
        
        return phones[0] if phones else "Not specified"

    def extract_name(self, text):
        lines = text.split('\n')
        
        # Look for name in first few lines, avoiding common headers
        skip_patterns = [
            r'resume|cv|curriculum vitae',
            r'@|\.com|\.org|\.edu',
            r'\d{3,}',  # Numbers (likely phone/address)
            r'street|road|avenue|city|state|country',
            r'phone|email|address|contact'
        ]
        
        for line in lines[:8]:  # Check more lines
            line = line.strip()
            if not line or len(line) < 2:
                continue
                
            # Skip if matches skip patterns
            if any(re.search(pattern, line.lower()) for pattern in skip_patterns):
                continue
            
            # Check if it looks like a name
            words = line.split()
            if 2 <= len(words) <= 4:  # Names typically 2-4 words
                # Check if words start with capital letters
                if all(word[0].isupper() for word in words if word):
                    return line
        
        return "Not specified"

    def extract_skills(self, text):
        found_skills = []
        text_lower = text.lower()
        
        # Look for skills in dedicated sections first
        skill_sections = re.findall(r'(?:skills?|technical skills?|competencies)[:\-\s]*(.*?)(?:\n\n|$)', 
                                   text, re.IGNORECASE | re.DOTALL)
        
        search_text = ' '.join(skill_sections) + ' ' + text_lower
        
        for skill in self.technical_skills:
            # Use word boundaries for better matching
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, search_text):
                found_skills.append(skill.title())
        
        # Remove duplicates while preserving order
        found_skills = list(dict.fromkeys(found_skills))
        return ', '.join(found_skills[:15]) if found_skills else "Not specified"  # Limit to 15 skills

    def extract_education(self, text):
        education_info = []
        lines = text.split('\n')
        
        # Look for education section first
        education_section = self._find_section(text, ['education', 'academic', 'qualification'])
        
        if education_section:
            search_lines = education_section.split('\n')
        else:
            search_lines = lines
        
        degree_patterns = [
            r'\b(bachelor|master|phd|doctorate|mba|ms|ma|bs|ba|btech|be|bsc|msc|bca|mca)\b.*?(\d{4}|\d{2})',
            r'\b(university|college|institute)\s+of\s+[\w\s]+',
            r'\b[\w\s]+(university|college|institute)\b',
        ]
        
        for line in search_lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if line contains education keywords
            if any(keyword in line.lower() for keyword in self.education_keywords):
                # Look for degree patterns
                for pattern in degree_patterns:
                    if re.search(pattern, line, re.IGNORECASE):
                        education_info.append(line)
                        break
                else:
                    # If no specific pattern, but contains education keywords
                    if len(line) > 10:
                        education_info.append(line)
        
        # Remove duplicates and clean
        education_info = list(dict.fromkeys(education_info))
        education_text = ' | '.join(education_info[:5])  # Limit to 5 entries
        
        return education_text[:800] if education_text else "Not specified"

    def extract_experience(self, text):
        # Look for experience section
        experience_section = self._find_section(text, ['experience', 'work', 'employment', 'career'])
        
        if experience_section:
            search_text = experience_section
        else:
            search_text = text
        
        experience_entries = []
        lines = search_text.split('\n')
        
        # Patterns for job titles and companies
        job_patterns = [
            r'\b(manager|developer|engineer|analyst|coordinator|specialist|associate|director|executive|intern)\b',
            r'\b(senior|junior|lead|principal|chief|head)\s+\w+',
            r'\b\w+\s+(manager|developer|engineer|analyst)\b'
        ]
        
        # Date patterns
        date_patterns = [
            r'\b(20\d{2}|19\d{2})\b',
            r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s*(20\d{2}|19\d{2}|\d{2})\b',
            r'\b\d{1,2}/\d{1,2}/\d{2,4}\b'
        ]
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            # Check if line contains job title or company info
            has_job_pattern = any(re.search(pattern, line, re.IGNORECASE) for pattern in job_patterns)
            has_date = any(re.search(pattern, line, re.IGNORECASE) for pattern in date_patterns)
            
            if has_job_pattern or has_date:
                # Collect this entry and next few lines for context
                entry_lines = [line]
                j = i + 1
                while j < len(lines) and j < i + 4:  # Get next 3 lines max
                    next_line = lines[j].strip()
                    if next_line and len(next_line) > 5:
                        entry_lines.append(next_line)
                    j += 1
                
                experience_entries.append(' | '.join(entry_lines))
                i = j
            else:
                i += 1
        
        # Remove duplicates and join
        experience_entries = list(dict.fromkeys(experience_entries))
        experience_text = ' || '.join(experience_entries[:8])  # Limit to 8 entries
        
        return experience_text[:1200] if experience_text else "Not specified"

    def extract_location(self, text):
        # Enhanced location patterns
        location_patterns = [
            r'([A-Za-z\s]+),\s*([A-Za-z\s]+)\s*,?\s*(\d{5,6}|\d{3}\s*\d{3})?',  # City, State, ZIP
            r'([A-Za-z\s]+),\s*([A-Za-z]{2,})',  # City, State/Country
            r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*([A-Z]{2}|\w+)\b',  # Formatted location
        ]
        
        # Look near contact information
        contact_section = text[:500]  # First part likely has contact info
        
        for pattern in location_patterns:
            matches = re.findall(pattern, contact_section)
            if matches:
                match = matches[0]
                if isinstance(match, tuple):
                    location = ', '.join([part for part in match if part.strip()])
                else:
                    location = match
                
                # Validate it looks like a real location
                if len(location) > 3 and not re.search(r'@|\.com|phone|email', location.lower()):
                    return location
        
        return "Not specified"

    def extract_summary(self, text):
        # Look for dedicated summary sections
        summary_section = self._find_section(text, ['summary', 'objective', 'profile', 'about'])
        
        if summary_section:
            # Clean and return the summary section
            summary = re.sub(r'\n+', ' ', summary_section)
            summary = re.sub(r'\s+', ' ', summary)
            return summary.strip()[:500]
        
        # If no dedicated section, look for paragraph-like content
        lines = text.split('\n')
        for i, line in enumerate(lines[:15]):  # Check first 15 lines
            line = line.strip()
            
            # Skip headers, contact info, etc.
            if (len(line) > 80 and 
                not re.search(r'@|\.com|\d{3,}|phone|email|address', line.lower()) and
                not line.isupper()):
                
                # This might be a summary paragraph
                return line[:500]
        
        return "Not specified"

    def extract_languages(self, text):
        found_languages = []
        text_lower = text.lower()
        
        # Look for language section
        language_section = self._find_section(text, ['language', 'languages'])
        search_text = language_section + ' ' + text_lower if language_section else text_lower
        
        for lang in self.language_keywords:
            if lang in search_text:
                if lang not in ['fluent', 'native', 'conversational', 'basic', 'intermediate', 'advanced']:
                    found_languages.append(lang.title())
        
        # Remove duplicates
        found_languages = list(dict.fromkeys(found_languages))
        return ', '.join(found_languages[:5]) if found_languages else "Not specified"

    def extract_certifications(self, text):
        found_certs = []
        text_lower = text.lower()
        
        # Look for certification section
        cert_section = self._find_section(text, ['certification', 'certifications', 'licenses'])
        search_text = cert_section + ' ' + text_lower if cert_section else text_lower
        
        for cert in self.certification_keywords:
            if cert in search_text:
                found_certs.append(cert.title())
        
        # Also look for specific certification patterns
        cert_patterns = [
            r'(aws|microsoft|google|oracle|cisco)\s+certified\s+[\w\s]+',
            r'pmp|csm|csa|cissp|comptia\s+[\w+]+',
            r'certified\s+[\w\s]+\s+(professional|associate|expert)'
        ]
        
        for pattern in cert_patterns:
            matches = re.findall(pattern, search_text, re.IGNORECASE)
            found_certs.extend(matches)
        
        # Remove duplicates and clean
        found_certs = list(dict.fromkeys(found_certs))
        return ', '.join(found_certs[:8]) if found_certs else "Not specified"

    def _find_section(self, text, keywords):
        """Find and extract specific sections from resume"""
        lines = text.split('\n')
        section_content = []
        in_section = False
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Check if this line is a section header
            if any(keyword in line_lower for keyword in keywords):
                in_section = True
                continue
            
            # If we're in the section, collect content
            if in_section:
                # Stop if we hit another section header
                other_sections = ['experience', 'education', 'skills', 'projects', 'awards', 'contact']
                if (line_lower in other_sections or 
                    (line.isupper() and len(line.strip()) > 3 and len(line.strip()) < 30)):
                    break
                
                if line.strip():
                    section_content.append(line)
                elif section_content:  # Empty line after content might end section
                    # Check next few lines to see if section continues
                    next_lines_empty = True
                    for j in range(i+1, min(i+3, len(lines))):
                        if lines[j].strip():
                            next_lines_empty = False
                            break
                    if next_lines_empty:
                        break
        
        return '\n'.join(section_content) if section_content else ""

    def parse_resume(self, text):
        """Enhanced main parsing function"""
        try:
            parsed_data = {
                'name': self.extract_name(text),
                'email': self.extract_email(text),
                'phone': self.extract_phone(text),
                'location': self.extract_location(text),
                'summary': self.extract_summary(text),
                'experience': self.extract_experience(text),
                'education': self.extract_education(text),
                'skills': self.extract_skills(text),
                'certifications': self.extract_certifications(text),
                'languages': self.extract_languages(text)
            }
            return parsed_data
        except Exception as e:
            st.error(f"Error parsing resume: {str(e)}")
            return None

# Clean text for Excel compatibility
def clean_text_for_excel(text):
    """Remove illegal characters that can't be used in Excel worksheets"""
    if not isinstance(text, str):
        return text
    
    # Remove control characters (ASCII 0-31 except tab, newline, carriage return)
    # Excel doesn't allow these characters
    illegal_chars = []
    for i in range(32):
        if i not in [9, 10, 13]:  # Keep tab, newline, carriage return
            illegal_chars.append(chr(i))
    
    # Also remove some other problematic characters
    illegal_chars.extend(['\x7f', '\x80', '\x81', '\x82', '\x83', '\x84', '\x85', '\x86', '\x87', 
                         '\x88', '\x89', '\x8a', '\x8b', '\x8c', '\x8d', '\x8e', '\x8f', '\x90', 
                         '\x91', '\x92', '\x93', '\x94', '\x95', '\x96', '\x97', '\x98', '\x99', 
                         '\x9a', '\x9b', '\x9c', '\x9d', '\x9e', '\x9f'])
    
    # Remove illegal characters
    for char in illegal_chars:
        text = text.replace(char, '')
    
    # Also clean up any remaining problematic unicode characters
    # Replace with similar safe characters
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    
    # Limit length to prevent Excel issues
    if len(text) > 32767:  # Excel cell character limit
        text = text[:32767]
    
    return text.strip()

# Clean DataFrame for Excel export
def clean_dataframe_for_excel(df):
    """Clean all text columns in DataFrame for Excel compatibility"""
    df_clean = df.copy()
    
    for column in df_clean.columns:
        if df_clean[column].dtype == 'object':  # Text columns
            df_clean[column] = df_clean[column].apply(clean_text_for_excel)
    
    return df_clean

# Initialize parser
parser = ResumeParser()

# Get all PDF and DOCX files from folder
def get_resume_files_from_folder(folder_path):
    if not os.path.exists(folder_path):
        return []
    
    pdf_files = glob.glob(os.path.join(folder_path, "*.pdf"))
    docx_files = glob.glob(os.path.join(folder_path, "*.docx"))
    
    return pdf_files + docx_files

# Extract text from PDF
def extract_text_from_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        st.error(f"Error extracting PDF text: {str(e)}")
        return None

# Extract text from DOCX (handles both file path and file object)
def extract_text_from_docx(file_path_or_file):
    try:
        if isinstance(file_path_or_file, str):
            doc = docx.Document(file_path_or_file)
        else:
            doc = docx.Document(file_path_or_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting DOCX text: {str(e)}")
        return None

# Read file content from path
def read_file_content(file_path):
    try:
        if file_path.lower().endswith('.pdf'):
            with open(file_path, 'rb') as file:
                return extract_text_from_pdf(file)
        elif file_path.lower().endswith('.docx'):
            return extract_text_from_docx(file_path)
    except Exception as e:
        st.error(f"Error reading {file_path}: {str(e)}")
        return None

# Streamlit App
def main():
    st.set_page_config(
        page_title="Fast Resume Parser",
        page_icon="⚡",
        layout="wide"
    )

    st.title("⚡ Fast Resume Parser (No AI Required)")
    st.markdown("Extract resume information using smart pattern matching - **Lightning Fast!**")

    # Sidebar
    with st.sidebar:
        
        
        st.markdown("---")
        st.header("Instructions")
        st.markdown("""
        **Folder Processing:**
        1. Enter folder path with CV files
        2. Click 'Process All Resumes'
        3. Download the Excel file
        
        **Manual Upload:**
        - Upload individual PDF/DOCX files
        - Instant processing
        """)
        
        st.markdown("---")
        st.header("Extracted Fields")
        st.markdown("""
        • Name & Contact Info
        • Email & Phone
        • Location
        • Professional Summary
        • Work Experience
        • Education
        • Skills
        • Certifications
        • Languages
        """)

    # Main content
    col1, col2 = st.columns([2, 1])

    with col1:
        st.header("Select Resume Folder")
        
        # Folder path input
        folder_path = st.text_input(
            "Enter folder path containing CV files:",
            placeholder="e.g., C:/Users/Documents/Resumes or /home/user/resumes",
            help="Enter the full path to the folder containing PDF and DOCX resume files"
        )
        
        # Alternative: Show current directory option
        if st.checkbox("Use current directory"):
            folder_path = os.getcwd()
            st.info(f"Using current directory: {folder_path}")
        
        # Show files in selected folder
        if folder_path:
            if os.path.exists(folder_path):
                resume_files = get_resume_files_from_folder(folder_path)
                
                if resume_files:
                    st.success(f"Found {len(resume_files)} resume files:")
                    
                    # Display found files
                    files_df = pd.DataFrame({
                        'Filename': [os.path.basename(f) for f in resume_files],
                        'Type': [os.path.splitext(f)[1].upper() for f in resume_files],
                        'Size (KB)': [round(os.path.getsize(f)/1024, 2) for f in resume_files]
                    })
                    st.dataframe(files_df, use_container_width=True)
                    
                    # Process button
                    if st.button("⚡ Process All Resumes (Fast!)", type="primary"):
                        # Progress bar
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        start_time = datetime.now()
                        
                        parsed_data = []
                        successful_parses = 0
                        failed_parses = 0
                        failed_files = []
                        
                        for i, file_path in enumerate(resume_files):
                            filename = os.path.basename(file_path)
                            status_text.text(f"Processing {filename}...")
                            
                            # Extract text
                            text = read_file_content(file_path)
                            
                            if text and text.strip():
                                # Parse with rule-based parser
                                parsed_info = parser.parse_resume(text)
                                
                                if parsed_info:
                                    parsed_info['filename'] = filename
                                    parsed_info['file_path'] = file_path
                                    parsed_info['parsed_date'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                                    parsed_data.append(parsed_info)
                                    successful_parses += 1
                                else:
                                    failed_parses += 1
                                    failed_files.append({'filename': filename, 'reason': 'Parsing failed'})
                            else:
                                failed_parses += 1
                                failed_files.append({'filename': filename, 'reason': 'Text extraction failed or empty file'})
                            
                            # Update progress
                            progress_bar.progress((i + 1) / len(resume_files))
                        
                        end_time = datetime.now()
                        processing_time = (end_time - start_time).total_seconds()
                        
                        status_text.text(f"Processing completed in {processing_time:.2f} seconds!")
                        
                        # Show results summary
                        col_success, col_failed, col_speed = st.columns(3)
                        with col_success:
                            st.metric("Successfully Parsed", successful_parses)
                        with col_failed:
                            st.metric("Failed to Parse", failed_parses)
                        with col_speed:
                            st.metric("Processing Time", f"{processing_time:.2f}s")
                        
                        # Display failed files if any
                        if failed_files:
                            st.error("❌ **Failed to Parse:**")
                            failed_df = pd.DataFrame(failed_files)
                            st.dataframe(
                                failed_df, 
                                use_container_width=True,
                                column_config={
                                    "filename": st.column_config.TextColumn("File Name"),
                                    "reason": st.column_config.TextColumn("Failure Reason")
                                }
                            )
                        
                        if parsed_data:
                            st.success(f"⚡ Successfully processed {len(parsed_data)} resumes in {processing_time:.2f} seconds!")
                            
                            if failed_parses > 0:
                                st.warning(f"⚠️ {failed_parses} files failed to parse. See error details above.")
                            
                            # Save to Excel
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            excel_filename = f"parsed_resumes_{timestamp}.xlsx"
                            
                            # Create DataFrame and save
                            df = pd.DataFrame(parsed_data)
                            df_export = df.drop('file_path', axis=1, errors='ignore')
                            
                            # Clean data for Excel compatibility
                            df_export = clean_dataframe_for_excel(df_export)
                            df_export.to_excel(excel_filename, index=False)
                            
                            # Display results
                            st.header("Parsed Results")
                            st.dataframe(df_export, use_container_width=True)
                            
                            # Download button
                            with open(excel_filename, "rb") as file:
                                st.download_button(
                                    label="📥 Download Excel File",
                                    data=file.read(),
                                    file_name=excel_filename,
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                )
                else:
                    st.warning("No PDF or DOCX files found in the specified folder")
            else:
                st.error("❌ Folder path does not exist. Please check the path and try again.")
        else:
            st.info("📁 Please enter a folder path to scan for resume files")

        # Manual file upload as backup option
        st.markdown("---")
        st.header("Alternative: Manual File Upload")
        uploaded_files = st.file_uploader(
            "Upload individual files for instant processing",
            type=['pdf', 'docx'],
            accept_multiple_files=True,
            help="Upload PDF or DOCX files - no API key needed!"
        )

        if uploaded_files:
            if st.button("⚡ Parse Uploaded Files (Instant!)", type="secondary"):
                # Progress bar
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                start_time = datetime.now()
                
                parsed_data = []
                successful_parses = 0
                failed_parses = 0
                failed_files = []
                
                for i, uploaded_file in enumerate(uploaded_files):
                    status_text.text(f"Processing {uploaded_file.name}...")
                    
                    # Extract text based on file type
                    if uploaded_file.type == "application/pdf":
                        text = extract_text_from_pdf(uploaded_file)
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        text = extract_text_from_docx(uploaded_file)
                    else:
                        failed_files.append({'filename': uploaded_file.name, 'reason': 'Unsupported file type'})
                        failed_parses += 1
                        continue
                    
                    if text and text.strip():
                        # Parse with rule-based parser
                        parsed_info = parser.parse_resume(text)
                        
                        if parsed_info:
                            parsed_info['filename'] = uploaded_file.name
                            parsed_info['parsed_date'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                            parsed_data.append(parsed_info)
                            successful_parses += 1
                        else:
                            failed_parses += 1
                            failed_files.append({'filename': uploaded_file.name, 'reason': 'Parsing failed'})
                    else:
                        failed_parses += 1
                        failed_files.append({'filename': uploaded_file.name, 'reason': 'Text extraction failed or empty file'})
                    
                    # Update progress
                    progress_bar.progress((i + 1) / len(uploaded_files))
                
                end_time = datetime.now()
                processing_time = (end_time - start_time).total_seconds()
                
                status_text.text(f"Processing completed in {processing_time:.2f} seconds!")
                
                # Show results summary
                col_success, col_failed, col_speed = st.columns(3)
                with col_success:
                    st.metric("Successfully Parsed", successful_parses)
                with col_failed:
                    st.metric("Failed to Parse", failed_parses)
                with col_speed:
                    st.metric("Processing Time", f"{processing_time:.2f}s")
                
                # Display failed files if any
                if failed_files:
                    st.error("❌ **Failed to Parse:**")
                    failed_df = pd.DataFrame(failed_files)
                    st.dataframe(
                        failed_df, 
                        use_container_width=True,
                        column_config={
                            "filename": st.column_config.TextColumn("File Name"),
                            "reason": st.column_config.TextColumn("Failure Reason")
                        }
                    )
                
                if parsed_data:
                    st.success(f"⚡ Successfully processed {len(parsed_data)} resumes in {processing_time:.2f} seconds!")
                    
                    if failed_parses > 0:
                        st.warning(f"⚠️ {failed_parses} files failed to parse. See error details above.")
                    
                    # Save to Excel
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    excel_filename = f"uploaded_resumes_{timestamp}.xlsx"
                    
                    df = pd.DataFrame(parsed_data)
                    
                    # Clean data for Excel compatibility
                    df_clean = clean_dataframe_for_excel(df)
                    df_clean.to_excel(excel_filename, index=False)
                    
                    # Display results
                    st.header("Parsed Results")
                    st.dataframe(df_clean, use_container_width=True)
                    
                    # Download button
                    with open(excel_filename, "rb") as file:
                        st.download_button(
                            label="📥 Download Excel File",
                            data=file.read(),
                            file_name=excel_filename,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )

    with col2:
        
        
      
        
        st.header("Accuracy Improvements")
        st.markdown("""
        🎯 **Enhanced Pattern Matching**  
        📊 **Section-Based Extraction**  
        🔍 **Smart Name Detection**  
        📞 **Better Phone/Email Parsing**  
        💼 **Job Title Recognition**  
        🎓 **Degree Pattern Matching**  
        🌐 **Language Detection**  
        📜 **Certification Extraction**  
        """)
        
        
        

if __name__ == "__main__":
    main()